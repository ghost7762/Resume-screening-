"""
Resume Parser Service
Extracts plain text from PDF, DOCX, and TXT files.
"""
import os
import re

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from a PDF file using PyPDF2."""
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 is not installed. Run: pip install PyPDF2")

    text = []
    try:
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    return clean_text(" ".join(text))


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    return clean_text("\n".join(paragraphs))


def extract_text_from_txt(filepath: str) -> str:
    """Extract text from a plain text file."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return clean_text(f.read())
    except Exception as e:
        raise ValueError(f"Failed to read text file: {str(e)}")


def parse_resume(filepath: str) -> str:
    """
    Auto-detect file type and extract text.
    Supports: .pdf, .docx, .doc, .txt
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        return extract_text_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.")


def clean_text(text: str) -> str:
    """Remove excessive whitespace and special characters."""
    # Replace multiple whitespace/newlines with single space
    text = re.sub(r"\s+", " ", text)
    # Remove non-printable characters
    text = re.sub(r"[^\x20-\x7E]", " ", text)
    return text.strip()
